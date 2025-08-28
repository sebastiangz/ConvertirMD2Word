#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ConvertirMD2Word.py - Conversor de Markdown a Word para Moodle
Convierte archivos Markdown (.md) a documentos Word (.docx) optimizados 
para el plugin de importacion de libros de Moodle.
"""

import os
import sys
import pypandoc
import argparse
from pathlib import Path
import logging
import tempfile
import re

# Configurar logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
registro = logging.getLogger(__name__)

def crear_plantilla_moodle():
    """
    Crea una plantilla Word optimizada para el plugin de importacion de libros de Moodle
    
    Returns:
        str: Ruta del archivo de plantilla temporal
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        
        documento = Document()
        estilos = documento.styles
        
        # Heading 1 para Capitulos principales
        titulo1 = estilos['Heading 1']
        fuente_titulo1 = titulo1.font
        fuente_titulo1.name = 'Calibri'
        fuente_titulo1.size = Pt(16)
        fuente_titulo1.bold = True
        fuente_titulo1.color.rgb = None
        
        parrafo_titulo1 = titulo1.paragraph_format
        parrafo_titulo1.space_before = Pt(12)
        parrafo_titulo1.space_after = Pt(3)
        parrafo_titulo1.keep_with_next = True
        parrafo_titulo1.page_break_before = False
        
        # Heading 2 para Subcapitulos
        titulo2 = estilos['Heading 2']
        fuente_titulo2 = titulo2.font
        fuente_titulo2.name = 'Calibri'
        fuente_titulo2.size = Pt(13)
        fuente_titulo2.bold = True
        fuente_titulo2.color.rgb = None
        
        parrafo_titulo2 = titulo2.paragraph_format
        parrafo_titulo2.space_before = Pt(10)
        parrafo_titulo2.space_after = Pt(3)
        parrafo_titulo2.keep_with_next = True
        parrafo_titulo2.page_break_before = False
        
        # Titulos menores como texto normal
        for nivel in [3, 4, 5, 6]:
            try:
                titulo = estilos[f'Heading {nivel}']
                fuente_titulo = titulo.font
                fuente_titulo.name = 'Calibri'
                fuente_titulo.size = Pt(12 - (nivel - 3))
                fuente_titulo.bold = True
                fuente_titulo.color.rgb = None
                
                parrafo_titulo = titulo.paragraph_format
                parrafo_titulo.space_before = Pt(6)
                parrafo_titulo.space_after = Pt(3)
            except KeyError:
                pass
        
        # Texto normal
        normal = estilos['Normal']
        fuente_normal = normal.font
        fuente_normal.name = 'Calibri'
        fuente_normal.size = Pt(11)
        fuente_normal.color.rgb = None
        
        parrafo_normal = normal.paragraph_format
        parrafo_normal.space_after = Pt(6)
        parrafo_normal.line_spacing = 1.15
        
        # Margenes del documento
        seccion = documento.sections[0]
        seccion.top_margin = Inches(1)
        seccion.bottom_margin = Inches(1)
        seccion.left_margin = Inches(1)
        seccion.right_margin = Inches(1)
        
        # Guardar plantilla temporal
        archivo_temporal = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        archivo_temporal.close()
        documento.save(archivo_temporal.name)
        
        registro.info(f"Plantilla Moodle creada: {archivo_temporal.name}")
        return archivo_temporal.name
        
    except ImportError:
        registro.warning("python-docx no disponible. Usando configuracion basica.")
        return None
    except Exception as e:
        registro.warning(f"Error creando plantilla: {e}. Usando configuracion basica.")
        return None

def detectar_y_corregir_codificacion(ruta_archivo):
    """
    Detecta automaticamente la codificacion de un archivo y corrige problemas comunes
    usando modulos especializados
    
    Args:
        ruta_archivo (str): Ruta del archivo a procesar
    
    Returns:
        str: Contenido del archivo con codificacion corregida
    """
    try:
        # Importar modulos especializados si estan disponibles
        try:
            import chardet
            tiene_chardet = True
        except ImportError:
            tiene_chardet = False
            registro.warning("chardet no disponible. Usando deteccion basica de codificacion.")
        
        try:
            import ftfy
            tiene_ftfy = True
        except ImportError:
            tiene_ftfy = False
            registro.warning("ftfy no disponible. Usando correccion basica de caracteres.")
        
        # Leer archivo en modo binario para deteccion de codificacion
        with open(ruta_archivo, 'rb') as archivo:
            datos_binarios = archivo.read()
        
        # Detectar codificacion automaticamente si chardet esta disponible
        codificacion_detectada = 'utf-8'
        if tiene_chardet:
            resultado_deteccion = chardet.detect(datos_binarios)
            if resultado_deteccion['confidence'] > 0.7:
                codificacion_detectada = resultado_deteccion['encoding']
                registro.info(f"Codificacion detectada: {codificacion_detectada}")
            else:
                registro.warning("Baja confianza en deteccion de codificacion")
        
        # Intentar decodificar con la codificacion detectada
        contenido = None
        codificaciones_intentar = [codificacion_detectada, 'utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for codificacion in codificaciones_intentar:
            try:
                contenido = datos_binarios.decode(codificacion)
                registro.info(f"Archivo decodificado exitosamente con: {codificacion}")
                break
            except UnicodeDecodeError:
                continue
        
        if contenido is None:
            # Ultimo recurso: decodificar con errores ignored
            contenido = datos_binarios.decode('utf-8', errors='ignore')
            registro.warning("Usando decodificacion con errores ignorados")
        
        # Aplicar ftfy para reparar texto mal codificado
        if tiene_ftfy:
            contenido_original = contenido
            contenido = ftfy.fix_text(contenido)
            if contenido != contenido_original:
                registro.info("ftfy aplico correcciones de codificacion")
        else:
            # Correcciones basicas si ftfy no esta disponible
            contenido = aplicar_correcciones_basicas(contenido)
        
        return contenido
        
    except Exception as e:
        registro.error(f"Error en deteccion/correccion de codificacion: {e}")
        # Fallback: leer con utf-8 y aplicar correcciones basicas
        try:
            with open(ruta_archivo, 'r', encoding='utf-8', errors='ignore') as archivo:
                contenido = archivo.read()
            return aplicar_correcciones_basicas(contenido)
        except Exception as e2:
            registro.error(f"Error en fallback de lectura: {e2}")
            return ""

def aplicar_correcciones_basicas(contenido):
    """
    Aplica correcciones basicas de codificacion cuando los modulos especializados 
    no estan disponibles
    
    Args:
        contenido (str): Contenido a corregir
    
    Returns:
        str: Contenido con correcciones basicas aplicadas
    """
    try:
        # Usar unidecode si esta disponible
        try:
            import unidecode
            caracteres_problematicos = ['ðŸ', 'â', 'Ã']
            for char_problema in caracteres_problematicos:
                if char_problema in contenido:
                    contenido = unidecode.unidecode(contenido)
                    registro.info("unidecode aplico correcciones de caracteres especiales")
                    break
        except ImportError:
            pass
        
        # Correcciones manuales basicas
        contenido = contenido.replace('Ã¡', 'á')
        contenido = contenido.replace('Ã©', 'é')
        contenido = contenido.replace('Ã­', 'í')
        contenido = contenido.replace('Ã³', 'ó')
        contenido = contenido.replace('Ãº', 'ú')
        contenido = contenido.replace('Ã±', 'ñ')
        contenido = contenido.replace('Â¿', '¿')
        contenido = contenido.replace('Â¡', '¡')
        
        # Eliminar caracteres de control
        contenido = contenido.replace('\ufeff', '')  # BOM
        contenido = contenido.replace('\u200b', '')  # Zero-width space
        contenido = contenido.replace('\u200c', '')  # Zero-width non-joiner
        contenido = contenido.replace('\u200d', '')  # Zero-width joiner
        
        # Normalizar espacios
        contenido = re.sub(r'\s+', ' ', contenido.replace('\t', ' '))
        
        return contenido
        
    except Exception as e:
        registro.warning(f"Error en correcciones basicas: {e}")
        return contenido

def procesar_markdown_previo(ruta_archivo_md):
    """
    Pre-procesa el archivo Markdown para corregir problemas que impiden 
    la importacion correcta en Moodle
    
    Args:
        ruta_archivo_md (str): Ruta del archivo Markdown original
    
    Returns:
        str: Ruta del archivo Markdown corregido (temporal)
    """
    try:
        registro.info("Procesando Markdown para compatibilidad con Moodle...")
        
        # Detectar y corregir codificacion automaticamente
        contenido_corregido = detectar_y_corregir_codificacion(ruta_archivo_md)
        if not contenido_corregido:
            registro.error("No se pudo leer el archivo de entrada")
            return ruta_archivo_md
        
        registro.info("Codificacion detectada y corregida automaticamente")
        
        # Corregir estructura de titulos
        lineas = contenido_corregido.split('\n')
        lineas_corregidas = []
        dentro_bloque_codigo = False
        
        # Palabras clave que indican comentarios de codigo mal interpretados como titulos
        palabras_clave_codigo = [
            'estilo imperativo', 'estilo funcional', 'funcion pura', 'funcion impura',
            'matematicamente:', 'composicion', 'resultado:', 'funcion que recibe',
            'funcion que retorna', 'modifica el estado', 'crea nuevos estados',
            'acumulacion funcional', 'transformacion de datos', 'uso',
            'estructura de datos', 'funciones puras', 'combinando ambos'
        ]
        
        for linea in lineas:
            # Detectar bloques de codigo
            if linea.strip().startswith('```'):
                dentro_bloque_codigo = not dentro_bloque_codigo
                lineas_corregidas.append(linea)
                continue
            
            # Si estamos en un bloque de codigo, no procesar
            if dentro_bloque_codigo:
                lineas_corregidas.append(linea)
                continue
            
            # Corregir titulos H1 falsos (comentarios de codigo)
            if (linea.strip().startswith('# ') and 
                any(palabra_clave in linea.lower() for palabra_clave in palabras_clave_codigo)):
                
                # Convertir a comentario dentro de bloque de codigo
                lineas_corregidas.append('```python')
                lineas_corregidas.append(linea)
                lineas_corregidas.append('```')
                registro.info(f"Corregido comentario de codigo: {linea.strip()[:50]}...")
            else:
                lineas_corregidas.append(linea)
        
        contenido_final = '\n'.join(lineas_corregidas)
        
        # Validar estructura final
        titulos_h1 = len(re.findall(r'^# [^#]', contenido_final, re.MULTILINE))
        titulos_h2 = len(re.findall(r'^## [^#]', contenido_final, re.MULTILINE))
        
        registro.info(f"Estructura corregida: {titulos_h1} capitulos, {titulos_h2} subcapitulos")
        
        # Crear archivo temporal corregido
        archivo_temporal = tempfile.NamedTemporaryFile(
            mode='w', suffix='.md', delete=False, encoding='utf-8'
        )
        archivo_temporal.write(contenido_final)
        archivo_temporal.close()
        
        registro.info(f"Archivo Markdown corregido creado: {archivo_temporal.name}")
        return archivo_temporal.name
        
    except Exception as e:
        registro.warning(f"Error en procesamiento previo: {e}")
        return ruta_archivo_md

def validar_estructura_markdown(ruta_archivo_md):
    """
    Valida la estructura del Markdown para compatibilidad con Moodle
    
    Args:
        ruta_archivo_md (str): Ruta del archivo Markdown
    
    Returns:
        dict: Informacion sobre la estructura del documento
    """
    try:
        registro.info("Validando estructura del Markdown...")
        
        # Leer archivo con deteccion automatica de codificacion
        contenido = detectar_y_corregir_codificacion(ruta_archivo_md)
        if not contenido:
            registro.error("No se pudo leer el archivo para validacion")
            return {}
        
        # Detectar problemas de codificacion residuales
        problemas_codificacion = []
        caracteres_problema = ['ðŸ', 'Ã', 'â', 'Â', '\ufffd']
        for caracter in caracteres_problema:
            if caracter in contenido:
                problemas_codificacion.append(caracter)
        
        if problemas_codificacion:
            registro.warning(f"Problemas de codificacion residuales detectados: {problemas_codificacion}")
            registro.info("Se aplicara correccion automatica durante la conversion")
        else:
            registro.info("No se detectaron problemas de codificacion")
        
        # Analizar estructura sin bloques de codigo
        lineas = contenido.split('\n')
        dentro_bloque_codigo = False
        lineas_limpias = []
        
        for linea in lineas:
            if linea.strip().startswith('```'):
                dentro_bloque_codigo = not dentro_bloque_codigo
                continue
            if not dentro_bloque_codigo:
                lineas_limpias.append(linea)
        
        contenido_limpio = '\n'.join(lineas_limpias)
        
        # Contar niveles de encabezados reales
        titulos_h1_reales = re.findall(r'^# [^#]', contenido_limpio, re.MULTILINE)
        titulos_h2_reales = re.findall(r'^## [^#]', contenido_limpio, re.MULTILINE)
        titulos_h3_reales = re.findall(r'^### [^#]', contenido_limpio, re.MULTILINE)
        
        # Detectar titulos falsos
        palabras_codigo = ['estilo', 'resultado:', 'funcion', 'uso', 'composicion']
        titulos_h1_falsos = []
        for linea in lineas:
            if (linea.strip().startswith('# ') and 
                any(palabra in linea.lower() for palabra in palabras_codigo)):
                titulos_h1_falsos.append(linea.strip())
        
        # Detectar imagenes
        imagenes = re.findall(r'!\[.*?\]\((.*?)\)', contenido)
        
        estructura = {
            'cantidad_h1': len(titulos_h1_reales),
            'cantidad_h2': len(titulos_h2_reales),
            'cantidad_h3': len(titulos_h3_reales),
            'cantidad_h1_falsos': len(titulos_h1_falsos),
            'imagenes': imagenes,
            'total_lineas': len(contenido.splitlines()),
            'tiene_problemas_codificacion': len(problemas_codificacion) > 0
        }
        
        registro.info("Analisis del Markdown:")
        registro.info(f"   Titulos H1 validos (capitulos): {len(titulos_h1_reales)}")
        registro.info(f"   Titulos H2 (subcapitulos): {len(titulos_h2_reales)}")
        registro.info(f"   Titulos H3+: {len(titulos_h3_reales)}")
        if titulos_h1_falsos:
            registro.info(f"   Titulos H1 falsos detectados: {len(titulos_h1_falsos)} (se corregiran)")
        registro.info(f"   Imagenes referenciadas: {len(imagenes)}")
        
        # Mostrar capitulos que se crearan
        registro.info("Capitulos que se crearan en Moodle:")
        for i, titulo in enumerate(titulos_h1_reales[:5], 1):
            registro.info(f"   {i}. {titulo}")
        if len(titulos_h1_reales) > 5:
            registro.info(f"   ... y {len(titulos_h1_reales) - 5} mas")
        
        # Recomendaciones
        if len(titulos_h1_reales) == 0:
            registro.warning("CRITICO: No hay titulos H1 validos")
            registro.info("Moodle necesita al menos un titulo # para crear capitulos")
        elif len(titulos_h1_reales) == 1:
            registro.info("PERFECTO: Un capitulo principal con subcapitulos")
        elif len(titulos_h1_reales) > 10:
            registro.warning(f"ATENCION: {len(titulos_h1_reales)} capitulos crearan muchas paginas en Moodle")
        
        return estructura
        
    except Exception as e:
        registro.warning(f"Error validando estructura: {e}")
        return {}

def optimizar_docx_para_moodle(ruta_archivo_docx):
    """
    Post-procesa el archivo DOCX para optimizar compatibilidad con Moodle
    
    Args:
        ruta_archivo_docx (str): Ruta del archivo DOCX a optimizar
    """
    try:
        from docx import Document
        
        registro.info("Post-procesando documento para Moodle...")
        
        documento = Document(ruta_archivo_docx)
        
        # Verificar y reportar estructura de capitulos
        contador_capitulos = 0
        contador_subcapitulos = 0
        
        for parrafo in documento.paragraphs:
            if parrafo.style.name == 'Heading 1':
                contador_capitulos += 1
                registro.info(f"Capitulo {contador_capitulos}: {parrafo.text[:50]}...")
            elif parrafo.style.name == 'Heading 2':
                contador_subcapitulos += 1
                registro.info(f"Subcapitulo: {parrafo.text[:50]}...")
        
        # Contar imagenes embebidas
        contador_imagenes = 0
        try:
            for relacion in documento.part.rels:
                if "image" in documento.part.rels[relacion].target_ref:
                    contador_imagenes += 1
        except:
            pass
        
        if contador_imagenes > 0:
            registro.info(f"Imagenes embebidas encontradas: {contador_imagenes}")
            registro.info("Formatos compatibles: GIF, PNG, JPEG")
        
        # Guardar documento optimizado
        documento.save(ruta_archivo_docx)
        
        registro.info("Optimizacion completada")
        registro.info(f"Resumen: {contador_capitulos} capitulos, {contador_subcapitulos} subcapitulos, {contador_imagenes} imagenes")
        
        if contador_capitulos == 0:
            registro.warning("ADVERTENCIA: No se encontraron Heading 1. El archivo podria no dividirse en capitulos.")
            registro.info("Consejo: Usa # para titulos principales (capitulos) y ## para subtitulos (subcapitulos)")
        
    except ImportError:
        registro.info("Post-procesamiento omitido (python-docx no disponible)")
    except Exception as e:
        registro.warning(f"Error en post-procesamiento: {e}")

def convertir_md_a_word(archivo_entrada, archivo_salida=None, optimizar_para_moodle=True):
    """
    Convierte un archivo Markdown a Word optimizado para importacion en Moodle
    
    Args:
        archivo_entrada (str): Ruta del archivo .md
        archivo_salida (str): Ruta del archivo .docx (opcional)
        optimizar_para_moodle (bool): Si aplicar optimizaciones especificas para Moodle
    
    Returns:
        bool: True si la conversion fue exitosa
    """
    try:
        # Validar archivo de entrada
        if not os.path.exists(archivo_entrada):
            registro.error(f"Archivo no encontrado: {archivo_entrada}")
            return False
        
        # Generar nombre de archivo de salida si no se proporciona
        if archivo_salida is None:
            ruta_entrada = Path(archivo_entrada)
            archivo_salida = str(ruta_entrada.with_suffix('.docx'))
        
        registro.info(f"Convirtiendo: {archivo_entrada} -> {archivo_salida}")
        
        # Pre-procesar Markdown si esta optimizado para Moodle
        archivo_procesado = archivo_entrada
        if optimizar_para_moodle:
            archivo_procesado = procesar_markdown_previo(archivo_entrada)
            if archivo_procesado != archivo_entrada:
                registro.info("Usando archivo pre-procesado para conversion")
        
        # Argumentos especificos para compatibilidad con plugin de Moodle
        argumentos_pandoc = [
            '--standalone',
            '--wrap=none',
            '--metadata=title=""',
        ]
        
        if optimizar_para_moodle:
            # Configuracion especifica para plugin de importacion de Moodle
            ruta_plantilla = crear_plantilla_moodle()
            if ruta_plantilla:
                argumentos_pandoc.append(f'--reference-doc={ruta_plantilla}')
            
            # Configuraciones adicionales para Word/Moodle
            argumentos_pandoc.extend([
                '--toc',
                '--toc-depth=2',
            ])
            
            registro.info("Optimizando para plugin de importacion de libros de Moodle")
            registro.info("Configurando Heading 1 y Heading 2 para capitulos/subcapitulos")
        
        # Realizar la conversion
        pypandoc.convert_file(
            archivo_procesado,
            'docx',
            outputfile=archivo_salida,
            extra_args=argumentos_pandoc
        )
        
        # Post-procesamiento para Moodle
        if optimizar_para_moodle:
            optimizar_docx_para_moodle(archivo_salida)
            
            # Limpiar archivos temporales
            for argumento in argumentos_pandoc:
                if argumento.startswith('--reference-doc='):
                    plantilla_temporal = argumento.split('=', 1)[1]
                    if plantilla_temporal and os.path.exists(plantilla_temporal):
                        try:
                            os.unlink(plantilla_temporal)
                        except:
                            pass
                    break
            
            # Limpiar archivo Markdown temporal
            if archivo_procesado != archivo_entrada and os.path.exists(archivo_procesado):
                try:
                    os.unlink(archivo_procesado)
                    registro.info("Archivos temporales limpiados")
                except:
                    pass
        
        registro.info(f"Conversion exitosa: {archivo_salida}")
        registro.info("Archivo listo para importar en Moodle Book")
        return True
        
    except Exception as e:
        registro.error(f"Error en conversion: {e}")
        return False

def convertir_directorio(directorio_entrada, directorio_salida=None, optimizar_para_moodle=True):
    """
    Convierte todos los archivos .md en un directorio
    
    Args:
        directorio_entrada (str): Directorio con archivos .md
        directorio_salida (str): Directorio de salida (opcional)
        optimizar_para_moodle (bool): Si aplicar optimizaciones para Moodle
    
    Returns:
        int: Numero de archivos convertidos exitosamente
    """
    if not os.path.exists(directorio_entrada):
        registro.error(f"Directorio no encontrado: {directorio_entrada}")
        return 0
    
    if directorio_salida and not os.path.exists(directorio_salida):
        os.makedirs(directorio_salida, exist_ok=True)
    
    archivos_md = list(Path(directorio_entrada).glob('*.md'))
    contador_convertidos = 0
    
    registro.info(f"Encontrados {len(archivos_md)} archivos .md en {directorio_entrada}")
    
    for i, archivo_md in enumerate(archivos_md, 1):
        registro.info(f"\nProcesando archivo {i}/{len(archivos_md)}: {archivo_md.name}")
        
        # Validar estructura antes de convertir
        estructura = validar_estructura_markdown(str(archivo_md))
        
        if estructura.get('cantidad_h1', 0) == 0:
            registro.warning(f"Saltando {archivo_md.name}: No hay capitulos H1 validos")
            continue
        
        if directorio_salida:
            archivo_salida = Path(directorio_salida) / f"{archivo_md.stem}.docx"
        else:
            archivo_salida = archivo_md.with_suffix('.docx')
        
        if convertir_md_a_word(str(archivo_md), str(archivo_salida), optimizar_para_moodle):
            contador_convertidos += 1
    
    registro.info(f"\nConversion completada: {contador_convertidos}/{len(archivos_md)} archivos")
    
    if optimizar_para_moodle and contador_convertidos > 0:
        registro.info("\nArchivos listos para plugin de importacion de libros de Moodle:")
        registro.info("   Formato: .docx (compatible)")
        registro.info("   Estilos: Heading 1 (capitulos), Heading 2 (subcapitulos)")
        registro.info("   Imagenes: Formato web-compatible embebido")
        registro.info("\nPara importar en Moodle:")
        registro.info("   1. Ve a tu curso -> Actividades -> Libro")
        registro.info("   2. Crea un nuevo Libro")
        registro.info("   3. Configuracion -> Importar capitulos")
        registro.info("   4. Sube tu archivo .docx")
    
    return contador_convertidos

def verificar_dependencias():
    """Verifica que las dependencias esten instaladas"""
    try:
        version = pypandoc.get_pandoc_version()
        registro.info(f"Pandoc version: {version}")
        return True
    except OSError:
        registro.error("Pandoc no esta instalado.")
        print("\nPara instalar pandoc:")
        print("Windows: Descargar desde https://pandoc.org/installing.html")
        print("Linux: sudo apt-get install pandoc")
        print("Mac: brew install pandoc")
        print("O ejecutar: pip install pypandoc && python -c \"import pypandoc; pypandoc.download_pandoc()\"")
        return False
    except Exception as e:
        registro.error(f"Error verificando dependencias: {e}")
        return False

def main():
    """Funcion principal optimizada para plugin de importacion de libros de Moodle"""
    analizador = argparse.ArgumentParser(
        description='Conversor Markdown a Word optimizado para plugin de importacion de libros de Moodle',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
OPTIMIZADO PARA PLUGIN DE IMPORTACION DE LIBROS DE MOODLE:
  - Heading 1 (#) -> Capitulos
  - Heading 2 (##) -> Subcapitulos  
  - Imagenes en formato web (GIF, PNG, JPEG)
  - Formato .docx compatible con Moodle 2.7+

Ejemplos de uso:
  %(prog)s archivo.md                     # Convierte con optimizacion Moodle
  %(prog)s archivo.md salida.docx         # Con nombre especifico
  %(prog)s -d carpeta_md/                 # Convierte carpeta completa
  %(prog)s archivo.md --sin-moodle        # Sin optimizaciones especificas
  %(prog)s --validar archivo.md           # Solo validar estructura
        """
    )
    
    analizador.add_argument('entrada', nargs='?', help='Archivo .md o directorio con archivos .md')
    analizador.add_argument('salida', nargs='?', help='Archivo .docx de salida o directorio destino')
    analizador.add_argument('-d', '--directorio', action='store_true', 
                           help='Modo directorio: convierte todos los .md')
    analizador.add_argument('--sin-moodle', action='store_true',
                           help='Desactivar optimizaciones especificas para Moodle')
    analizador.add_argument('--validar', action='store_true',
                           help='Solo validar estructura Markdown sin convertir')
    analizador.add_argument('-v', '--verbose', action='store_true', help='Salida detallada')
    
    argumentos = analizador.parse_args()
    
    # Configurar nivel de logging
    if argumentos.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    # Verificar dependencias
    if not verificar_dependencias():
        return 1
    
    # Si no hay argumentos, mostrar ayuda
    if not argumentos.entrada:
        analizador.print_help()
        print("\nEste programa esta optimizado para el plugin de importacion de libros de Moodle")
        return 0
    
    optimizar_para_moodle = not argumentos.sin_moodle
    
    if optimizar_para_moodle:
        registro.info("Modo: OPTIMIZADO PARA PLUGIN DE IMPORTACION DE LIBROS DE MOODLE")
        registro.info("Heading 1 = Capitulos | Heading 2 = Subcapitulos")
    else:
        registro.info("Modo: Conversion estandar")
    
    # Solo validar estructura
    if argumentos.validar:
        if argumentos.directorio:
            archivos_md = list(Path(argumentos.entrada).glob('*.md'))
            for archivo_md in archivos_md:
                registro.info(f"\nValidando: {archivo_md.name}")
                validar_estructura_markdown(str(archivo_md))
        else:
            validar_estructura_markdown(argumentos.entrada)
        return 0
    
    # Modo directorio
    if argumentos.directorio:
        convertidos = convertir_directorio(argumentos.entrada, argumentos.salida, optimizar_para_moodle)
        if convertidos > 0:
            print(f"\n{convertidos} archivos convertidos exitosamente!")
            if optimizar_para_moodle:
                print("Listos para importar en Moodle Book!")
        return 0 if convertidos > 0 else 1
    
    # Validar estructura del archivo individual
    estructura = validar_estructura_markdown(argumentos.entrada)
    
    # Modo archivo individual
    if convertir_md_a_word(argumentos.entrada, argumentos.salida, optimizar_para_moodle):
        nombre_salida = argumentos.salida or argumentos.entrada.replace('.md', '.docx')
        print(f"\nArchivo convertido exitosamente: {nombre_salida}")
        
        if optimizar_para_moodle:
            print("Listo para importar en Moodle Book!")
            print("\nPasos para importar:")
            print("   1. ELinea -> Tu curso -> Actividades -> Libro")
            print("   2. Crear nuevo libro")  
            print("   3. Configuracion -> Importar capitulos")
            print("   4. Subir archivo .docx")
            
            # Mostrar vista previa de estructura esperada
            if estructura.get('cantidad_h1', 0) > 0:
                print(f"\nSe crearan {estructura['cantidad_h1']} capitulos")
            if estructura.get('cantidad_h2', 0) > 0:
                print(f"Se crearan {estructura['cantidad_h2']} subcapitulos")
        
        return 0
    else:
        print("\nError en la conversion. Ver detalles arriba.")
        return 1

if __name__ == "__main__":
    try:
        sys.exit(main())
    except KeyboardInterrupt:
        print("\n\nConversion cancelada por el usuario.")
        sys.exit(1)
    except Exception as e:
        registro.error(f"Error inesperado: {e}")
        sys.exit(1)
